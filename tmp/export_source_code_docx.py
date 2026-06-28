from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


PROJECT_ROOT = Path("/Users/houqing/本科/学校作业/算法设计与分析/算法设计与分析project")
OUTPUT_PATH = PROJECT_ROOT / "output" / "docx" / "StudySmartAI_Source_Code.docx"
SOURCE_FOLDERS = ["include", "src"]


def set_run_font(run, font_name: str, font_size: int) -> None:
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def build_document() -> Document:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)
    normal_style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal_style.paragraph_format.space_after = Pt(8)
    normal_style.paragraph_format.line_spacing = 1.25

    code_style = document.styles.add_style("CodeLine", WD_STYLE_TYPE.PARAGRAPH)
    code_style.base_style = normal_style
    code_style.paragraph_format.space_after = Pt(0)
    code_style.paragraph_format.line_spacing = 1.0

    title = document.add_paragraph()
    title.style = document.styles["Title"]
    title_run = title.add_run("StudySmart AI Source Code")
    set_run_font(title_run, "Calibri", 24)

    for folder_name in SOURCE_FOLDERS:
        document.add_heading(folder_name, level=1)
        folder_path = PROJECT_ROOT / folder_name

        for source_path in sorted(folder_path.iterdir()):
            if source_path.suffix not in {".h", ".cpp"}:
                continue

            document.add_heading(source_path.name, level=2)
            for line in source_path.read_text(encoding="utf-8").splitlines():
                paragraph = document.add_paragraph(style=code_style)
                run = paragraph.add_run(line if line else " ")
                set_run_font(run, "Courier New", 9)

    return document


def main() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
